from __future__ import annotations

import io

from shared.types import ExtractionError, ExtractionResult


def extract_docx(data: bytes) -> ExtractionResult:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(data))

    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise ExtractionError("empty_content", "DOCX has no extractable text")

    full_text = "\n\n".join(parts)
    return ExtractionResult(
        text=full_text,
        metadata={
            "word_count": len(full_text.split()),
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    )
